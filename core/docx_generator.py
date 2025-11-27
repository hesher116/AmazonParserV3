"""DOCX Generator - Creates Word documents from parsed data"""
import glob
import os
import re
from typing import Dict, List
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from utils.logger import get_logger

logger = get_logger(__name__)


class DocxGenerator:
    """Generates Word documents from parsed Amazon product data."""
    
    def __init__(self):
        self.doc = None
    
    def generate(self, results: Dict, output_path: str) -> bool:
        """
        Generate DOCX document from results.
        
        Args:
            results: Dictionary with all parsed data
            output_path: Path to save the document
            
        Returns:
            True if successful
        """
        try:
            self.doc = Document()
            self._setup_styles()
            
            text_data = results.get('text', {})
            reviews_data = results.get('reviews', {})
            qa_data = results.get('qa', {})
            
            # Get product URL from results
            product_url = results.get('url', '')
            
            # 1. Title
            product_title = text_data.get('title', 'Unknown Product')
            if not product_title or product_title == 'Unknown Product':
                # Try to get from output directory name
                output_dir = results.get('output_dir', '')
                if output_dir:
                    from pathlib import Path
                    dir_name = Path(output_dir).name
                    # Remove numerical suffix like (2), (3)
                    import re
                    product_title = re.sub(r'\s*\(\d+\)\s*$', '', dir_name)
            
            self._add_title(product_title)
            
            # 2. URL (right after title)
            if product_url:
                self._add_url(product_url)
            
            # 3. Basic Info (Brand, ASIN, Price)
            self._add_basic_info(text_data)
            
            # 4. Product Overview
            if text_data.get('product_overview'):
                self._add_section('Product Overview', text_data['product_overview'])
            
            # 5. About This Item
            if text_data.get('about_this_item'):
                self._add_bullet_list('About This Item', text_data['about_this_item'])
            
            # 6. Ingredients
            if text_data.get('ingredients'):
                self._add_text_section('Ingredients', text_data['ingredients'])
            
            # 7. Important Information
            if text_data.get('important_information'):
                self._add_section('Important Information', text_data['important_information'])
            
            # 8. Technical Details
            if text_data.get('technical_details'):
                self._add_section('Technical Details', text_data['technical_details'])
            
            # 9. Product Details
            if text_data.get('product_details'):
                self._add_section('Product Details', text_data['product_details'])
            
            # 10. Customer Reviews Summary
            if reviews_data.get('summary'):
                self._add_reviews_summary(reviews_data['summary'])
            
            # 11. Review Details
            if reviews_data.get('reviews'):
                self._add_review_details(reviews_data['reviews'])
            
            # 12. Q&A
            if qa_data.get('qa_pairs'):
                self._add_qa_section(qa_data['qa_pairs'])
            
            # 13. Add images from folders (in correct order)
            output_dir = results.get('output_dir')
            if output_dir:
                self._add_images_from_folders(output_dir)
            
            # Save document
            self.doc.save(output_path)
            logger.info(f"Document saved: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to generate DOCX: {e}")
            return False
    
    def _setup_styles(self):
        """Setup document styles."""
        # Title style
        style = self.doc.styles['Title']
        style.font.size = Pt(24)
        style.font.bold = True
        
        # Heading 1
        style = self.doc.styles['Heading 1']
        style.font.size = Pt(16)
        style.font.bold = True
        
        # Heading 2
        style = self.doc.styles['Heading 2']
        style.font.size = Pt(14)
        style.font.bold = True
    
    def _add_title(self, title: str):
        """Add document title."""
        self.doc.add_heading(title, 0)
        self.doc.add_paragraph()
    
    def _add_url(self, url: str):
        """Add product URL after title."""
        p = self.doc.add_paragraph()
        p.add_run("URL: ").bold = True
        p.add_run(url)
        self.doc.add_paragraph()
    
    def _add_basic_info(self, text_data: Dict):
        """Add basic product information."""
        self.doc.add_heading('Product Information', level=1)
        
        info_items = []
        
        if text_data.get('brand'):
            info_items.append(f"Brand: {text_data['brand']}")
        
        if text_data.get('asin'):
            info_items.append(f"ASIN: {text_data['asin']}")
        
        price_data = text_data.get('price', {})
        if price_data.get('current_price'):
            price_str = f"Price: {price_data['current_price']}"
            if price_data.get('original_price'):
                price_str += f" (was {price_data['original_price']})"
            if price_data.get('savings'):
                price_str += f" - {price_data['savings']}"
            info_items.append(price_str)
        
        for item in info_items:
            self.doc.add_paragraph(item)
        
        self.doc.add_paragraph()
    
    def _add_section(self, title: str, data: Dict):
        """Add a section with key-value pairs."""
        self.doc.add_heading(title, level=1)
        
        # Create table for structured data
        if data:
            table = self.doc.add_table(rows=len(data), cols=2)
            table.style = 'Table Grid'
            
            for i, (key, value) in enumerate(data.items()):
                row = table.rows[i]
                row.cells[0].text = str(key)
                row.cells[1].text = str(value)
        
        self.doc.add_paragraph()
    
    def _add_bullet_list(self, title: str, items: List[str]):
        """Add a section with bullet points."""
        self.doc.add_heading(title, level=1)
        
        for item in items:
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(item)
        
        self.doc.add_paragraph()
    
    def _add_text_section(self, title: str, text: str):
        """Add a section with plain text."""
        self.doc.add_heading(title, level=1)
        self.doc.add_paragraph(text)
        self.doc.add_paragraph()
    
    def _add_reviews_summary(self, summary: Dict):
        """Add customer reviews summary section."""
        self.doc.add_heading('Customer Reviews', level=1)
        
        # Rating
        if summary.get('rating'):
            rating_text = f"Rating: {summary['rating']} out of 5 stars"
            if summary.get('rating_count'):
                rating_text += f" ({summary['rating_count']} ratings)"
            self.doc.add_paragraph(rating_text)
        
        # Star distribution
        if summary.get('star_distribution'):
            self.doc.add_heading('Rating Distribution', level=2)
            for stars, percent in sorted(summary['star_distribution'].items(), reverse=True):
                self.doc.add_paragraph(f"{stars}: {percent}")
        
        # Customers say
        if summary.get('customers_say'):
            self.doc.add_heading('Customers Say', level=2)
            self.doc.add_paragraph(summary['customers_say'])
        
        # Key aspects
        if summary.get('key_aspects'):
            self.doc.add_heading('Key Aspects', level=2)
            for aspect in summary['key_aspects']:
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(aspect)
        
        self.doc.add_paragraph()
    
    def _add_review_details(self, reviews: List[Dict]):
        """Add detailed reviews section."""
        self.doc.add_heading('Customer Review Details', level=1)
        
        for i, review in enumerate(reviews, 1):
            # Review header
            header = f"Review #{i}"
            if review.get('rating'):
                header += f" - {review['rating']}/5 stars"
            self.doc.add_heading(header, level=2)
            
            # Reviewer info
            info_parts = []
            if review.get('reviewer_name'):
                info_parts.append(f"By: {review['reviewer_name']}")
            if review.get('date'):
                info_parts.append(review['date'])
            if review.get('verified_purchase'):
                info_parts.append("Verified Purchase")
            if review.get('variant'):
                info_parts.append(review['variant'])
            
            if info_parts:
                self.doc.add_paragraph(' | '.join(info_parts))
            
            # Title
            if review.get('title'):
                p = self.doc.add_paragraph()
                p.add_run(review['title']).bold = True
            
            # Text
            if review.get('text'):
                self.doc.add_paragraph(review['text'])
            
            # Helpful count
            if review.get('helpful_count'):
                self.doc.add_paragraph(f"{review['helpful_count']} people found this helpful")
            
            self.doc.add_paragraph()
    
    def _add_qa_section(self, qa_pairs: List[Dict]):
        """Add Q&A section."""
        self.doc.add_heading('Questions & Answers', level=1)
        
        for i, qa in enumerate(qa_pairs, 1):
            # Question
            if qa.get('question'):
                p = self.doc.add_paragraph()
                p.add_run(f"Q{i}: ").bold = True
                p.add_run(qa['question'])
            
            # Answer
            if qa.get('answer'):
                p = self.doc.add_paragraph()
                p.add_run("A: ").bold = True
                p.add_run(qa['answer'])
            
            # Additional info
            info_parts = []
            if qa.get('votes'):
                info_parts.append(f"Votes: {qa['votes']}")
            if qa.get('answer_by'):
                info_parts.append(qa['answer_by'])
            
            if info_parts:
                p = self.doc.add_paragraph()
                p.add_run(' | '.join(info_parts)).italic = True
            
            self.doc.add_paragraph()
    
    def _add_images_from_folders(self, output_dir: str):
        """Add images from output folders to document in correct order."""
        base_path = Path(output_dir)
        
        if not base_path.exists():
            logger.warning(f"Output directory not found: {output_dir}")
            return
        
        # Define folders and their labels (in order of appearance)
        image_folders = [
            ('hero', 'Hero Image'),
            ('product', 'Product Gallery Images'),
            ('aplus_brand', 'A+ Content - From the Brand'),
            ('aplus_product', 'A+ Content - Product Description'),
            ('QAImages', 'Review Images'),
        ]
        
        # Collect all images with their order
        all_images = []
        
        for folder_name, section_title in image_folders:
            folder_path = base_path / folder_name
            
            if not folder_path.exists():
                continue
            
            # Find all images in folder
            img_files = glob.glob(str(folder_path / "*.*"))
            img_files = [f for f in img_files if f.lower().endswith(('.png', '.jpg', '.jpeg', '.webp'))]
            
            if not img_files:
                continue
            
            # Sort files to ensure correct numerical order (1, 2, 3, ..., 10, 11, not 1, 10, 11, 2, 3)
            def natural_sort_key(filename):
                """Natural sort key for filenames - sort by number only."""
                path = Path(filename)
                name = path.stem
                # Extract number if present - use only the number for sorting
                numbers = re.findall(r'\d+', name)
                if numbers:
                    # Return tuple: (0, number) to ensure numbers come first, sorted numerically
                    return (0, int(numbers[-1]))  # Use last number in case of multiple
                # If no number, sort alphabetically
                return (1, name)
            
            img_files.sort(key=natural_sort_key)
            
            # Add to all_images list with section info
            for img_path in img_files:
                all_images.append({
                    'path': img_path,
                    'section': section_title,
                    'folder': folder_name
                })
        
        # Add images in order, grouping by section
        current_section = None
        for img_info in all_images:
            # Add section heading if changed
            if img_info['section'] != current_section:
                if current_section is not None:
                    self.doc.add_paragraph()  # Space between sections
                self.doc.add_heading(img_info['section'], level=1)
                current_section = img_info['section']
            
            # Add image
            try:
                img_path = img_info['path']
                img_path_obj = Path(img_path)
                
                # Verify file exists and get info
                if not img_path_obj.exists():
                    raise FileNotFoundError(f"Image file not found: {img_path}")
                
                file_size = img_path_obj.stat().st_size
                if file_size == 0:
                    raise ValueError(f"Image file is empty: {img_path}")
                
                logger.debug(f"Adding image: {img_path_obj.name} ({file_size / 1024:.2f}KB)")
                
                # Try to add image - python-docx can have issues with certain image formats
                try:
                    # Use absolute path to avoid any path issues
                    abs_path = str(img_path_obj.resolve())
                    self.doc.add_picture(abs_path, width=Inches(4.0))
                except Exception as pic_error:
                    # Try with smaller size if first attempt fails
                    logger.debug(f"First attempt failed ({type(pic_error).__name__}), trying smaller size: {pic_error}")
                    try:
                        abs_path = str(img_path_obj.resolve())
                        self.doc.add_picture(abs_path, width=Inches(2.0))
                        logger.debug(f"Added with smaller size")
                    except Exception as pic_error2:
                        # Last resort: try to load with PIL and convert if needed
                        logger.debug(f"Second attempt failed, trying PIL conversion: {pic_error2}")
                        try:
                            from PIL import Image
                            from io import BytesIO
                            import tempfile
                            
                            # Open and verify image
                            img = Image.open(abs_path)
                            # Convert to RGB if needed (for JPEG compatibility)
                            if img.mode in ('RGBA', 'LA', 'P'):
                                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                                if img.mode == 'P':
                                    img = img.convert('RGBA')
                                rgb_img.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                                img = rgb_img
                            
                            # Save to temporary file
                            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as tmp_file:
                                img.save(tmp_file.name, 'JPEG', quality=95)
                                tmp_path = tmp_file.name
                            
                            # Add converted image
                            self.doc.add_picture(tmp_path, width=Inches(4.0))
                            
                            # Clean up temp file
                            try:
                                os.unlink(tmp_path)
                            except:
                                pass
                            
                            logger.debug(f"Added after PIL conversion")
                        except Exception as pic_error3:
                            raise pic_error3
                
                # Add filename as caption
                filename = img_path_obj.name
                p = self.doc.add_paragraph()
                p.add_run(filename).italic = True
                
                logger.debug(f"✓ Added image: {filename}")
                
            except Exception as e:
                # If image can't be loaded, add placeholder with detailed error
                filename = os.path.basename(img_info['path'])
                error_msg = str(e) if e else "Unknown error"
                error_type = type(e).__name__
                file_exists = Path(img_path).exists() if 'img_path' in locals() else False
                file_size = Path(img_path).stat().st_size if file_exists else 0
                
                self.doc.add_paragraph(f"[Image could not be loaded: {filename} - {error_type}: {error_msg}]")
                logger.warning(f"Failed to add image {filename}: {error_type}: {error_msg}")
                logger.debug(f"  Image path: {img_path if 'img_path' in locals() else 'N/A'}, exists: {file_exists}, size: {file_size} bytes")
        
        if all_images:
            self.doc.add_paragraph()

